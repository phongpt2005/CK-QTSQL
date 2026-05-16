import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

def add_heading_styled(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_table_from_rows(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()

def add_code_block(code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = bold
    return p

# Read the markdown file
with open(r'c:\wms\md\BAO_CAO_CSDL_TONG_HOP.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code = False
code_buf = []
in_table = False
table_headers = []
table_rows = []

def flush_table():
    global table_headers, table_rows, in_table
    if table_headers and table_rows:
        add_table_from_rows(table_headers, table_rows)
    table_headers = []
    table_rows = []
    in_table = False

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Code block
    if line.startswith('```'):
        if in_code:
            add_code_block('\n'.join(code_buf))
            code_buf = []
            in_code = False
        else:
            if in_table:
                flush_table()
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # Table row
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            # Check if separator row
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
        i += 1
        continue
    else:
        if in_table:
            flush_table()

    # Empty line
    if not line.strip():
        i += 1
        continue

    # Horizontal rule
    if line.strip() == '---':
        i += 1
        continue

    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        add_heading_styled(line[2:].strip(), 1)
        i += 1
        continue
    if line.startswith('## '):
        add_heading_styled(line[3:].strip(), 2)
        i += 1
        continue
    if line.startswith('### '):
        add_heading_styled(line[4:].strip(), 3)
        i += 1
        continue

    # Blockquote
    if line.startswith('>'):
        text = line.lstrip('> ').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        i += 1
        continue

    # Bullet points
    if line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        # Remove markdown bold
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.*)', line)
    if m:
        text = m.group(2).strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        i += 1
        continue

    # Normal paragraph
    text = line.strip()
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    if text:
        add_para(text)
    i += 1

# Flush remaining table
if in_table:
    flush_table()

output = r'c:\wms\md\BAO_CAO_CSDL_TONG_HOP.docx'
doc.save(output)
print(f'SUCCESS: File saved to {output}')
